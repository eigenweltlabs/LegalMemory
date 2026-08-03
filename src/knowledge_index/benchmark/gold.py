"""Gold generation: what a lawyer asks the assistant, machine-verified.

Gold is derived from the **source corpus files** (``mock_dms/`` + ``scenarios.jsonl``),
never from the database — so a benchmark run exercises the whole chain (convert,
extract, classify, version-chain, ACL, index, retrieve): any insertion failure
surfaces as a retrieval miss, and gold can never inherit an extraction mistake as
ground truth.

**One kind: a request.** This product is consumed through an assistant, not a search
box — the user speaks and the *agent* decides what to search for. An earlier version
also generated hand-written "search strings" as a second kind; 86% of them came back
byte-identical to their own answer, which asks a retriever to recompute the label's
definition (a lexical baseline satisfies it by construction) and tests a query
distribution that exists nowhere in the product. Reporting slices instead on
``meta.anchor`` — whether the request carries an identifier, an entity name, an
amount, or nothing — which preserves every ablation question worth asking.

The LLM only *proposes* ``(query, answer, source document)``; four mechanical checks
decide what becomes gold, with no human review step:

1. **No self-quoting** — a request that contains its own answer is not a question.
2. **Verbatim verification** — the answer must appear as a substring of a source
   document in the scenario's bundle, or the proposal is dropped (hallucinated gold
   is impossible). The model's document guess is a hint, never trusted.
3. **Corpus-wide discrimination** — the answer is then searched across *every*
   document the request's principal can see. A value found in more than
   ``max_gold_docs`` scope documents identifies nothing and is dropped.
4. **Dedupe** — by normalized query, so re-runs are safe and incremental.

Gold is **graded**: ``meta.primary_path`` is the document the answer was drawn from
(the one that answers the request); ``meta.secondary_paths`` merely repeat the value.
Flat labelling made a system that returned the answering document at rank 1 score
~0.64 instead of 1.0, and paid a bonus for retrieving incidental co-mentions.

Rejected proposals are logged with a diagnosed reason (``rejected.jsonl``) — nothing
is silently discarded. Everything runs through the LiteLLM gateway, so an air-gapped
install generates gold with its local model; only the gateway is needed — no
database, no index.
"""

from __future__ import annotations

import json
import random
import re
import threading
from collections import Counter, defaultdict
from concurrent.futures import ThreadPoolExecutor
from dataclasses import asdict, dataclass, field
from email import message_from_bytes
from pathlib import Path
from typing import Callable

from pydantic import BaseModel, Field

from knowledge_index.benchmark import gateway
from knowledge_index.config import AppConfig, ModelSlot

PROMPT_VERSION = "request-gold-3"
#: One kind — a request a lawyer makes of the assistant. Slicing is by ``anchor``
#: (below), not by kind: hand-written "search strings" tested a query distribution
#: that does not exist in this product, since the agent, not the user, forms queries.
KINDS: tuple[str, ...] = ("request",)

#: What hard handle (if any) the request carries — the axis the identifier and
#: lexical legs are supposed to serve. Derived from the request text, not invented.
_IDENTIFIER_SHAPE = re.compile(r"[A-Za-z]{2,}[-/]\d|\d{4,}|\b\d+[-/]\d+\b|§|\bNo\.\s*\d")
_ENTITY_SHAPE = re.compile(
    r"\b(LLP|LLC|L\.?P\.?|Inc\.?|Ltd\.?|N\.?A\.?|PLC|GmbH|AG|Corp\.?)\b", re.IGNORECASE
)
_AMOUNT_SHAPE = re.compile(r"[$€£]\s?\d|\b\d+(?:[.,]\d+)*\s*(?:%|percent|million|billion)\b")


def classify_anchor(query: str, answer: str = "") -> str:
    """Which hard handle the request turns on: identifier / entity / amount / none.

    Classified over the request AND its verified answer, because a request can turn
    on an identifier without containing one — "what's the docket number on the Apex
    complaint?" is an identifier lookup whose identifier lives in the answer. Judging
    the request text alone put those in the "none" bucket and left the identifier
    slice too small to compare, which is precisely the slice that decides whether the
    identifier leg earns its weight.
    """
    text = f"{query} {answer}"
    if _IDENTIFIER_SHAPE.search(text):
        return "identifier"
    if _AMOUNT_SHAPE.search(text):
        return "amount"
    if _ENTITY_SHAPE.search(text):
        return "entity"
    return "none"


_MAX_DOC_CHARS = 4000  # per-document text budget in the prompt
_WS = re.compile(r"\s+")


@dataclass
class GoldQuery:
    id: str
    kind: str
    query: str
    gold_paths: list[str]
    principals: list[str]
    matter_ref: str
    practice_area: str
    meta: dict = field(default_factory=dict)


def _normalize(text: str) -> str:
    return _WS.sub(" ", text).strip().casefold()


def _document_text(path: Path) -> str | None:
    """Raw text of a source file, for the kinds the grounder can read.

    ``.xlsx``/``.pdf``/binaries are not text-mined here; answers proposed from them
    fail verification and are diagnosed as ``answer_in_unreadable_file``.
    """
    suffix = path.suffix.casefold()
    try:
        if suffix == ".docx":
            from docx import Document as WordDocument

            document = WordDocument(str(path))
            parts = [p.text for p in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    parts.extend(cell.text for cell in row.cells)
            return "\n".join(parts)
        if suffix == ".eml":
            message = message_from_bytes(path.read_bytes())
            if message.is_multipart():
                chunks = [
                    part.get_content()
                    for part in message.walk()
                    if part.get_content_type() == "text/plain"
                ]
                return "\n".join(chunks)
            return message.get_content()
        if suffix in {".txt", ".md", ".csv"}:
            return path.read_text(encoding="utf-8", errors="replace")
    except Exception:
        return None
    return None


# ------------------------------------------------------------------ corpus text scope


def load_scenarios(corpus_dir: str | Path) -> list[dict]:
    return [
        json.loads(line)
        for line in (Path(corpus_dir) / "scenarios.jsonl").read_text(encoding="utf-8").splitlines()
        if line.strip()
    ]


def corpus_texts(corpus_dir: str | Path, scenarios: list[dict]) -> dict[str, tuple[str, str]]:
    """path → (principal, normalized text) for every readable document in the corpus.

    This is the discrimination scope: an answer's gold is every document *the query's
    principal can see* that states it — bundle-local verification alone would let a
    value that recurs across half the practice area masquerade as a known-item label.
    """
    source_root = Path(corpus_dir) / "mock_dms"
    texts: dict[str, tuple[str, str]] = {}
    for scenario in scenarios:
        for path in scenario["document_paths"]:
            if path in texts:
                continue
            raw = _document_text(source_root / path)
            if raw is not None and (normalized := _normalize(raw)):
                texts[path] = (scenario["principal"], normalized)
    return texts


def stratified_sample(scenarios: list[dict], limit: int | None, *, seed: int = 42) -> list[dict]:
    """Sample scenarios evenly across practice areas (seeded, deterministic).

    Round-robin over per-area shuffled lists, so a small ``limit`` still touches
    every practice area instead of whatever happens to sort first.
    """
    if limit is None or limit >= len(scenarios):
        return scenarios
    rng = random.Random(seed)
    by_area: dict[str, list[dict]] = defaultdict(list)
    for scenario in scenarios:
        by_area[scenario.get("practice_area", "")].append(scenario)
    for group in by_area.values():
        rng.shuffle(group)
    picked: list[dict] = []
    areas = sorted(by_area)
    index = 0
    while len(picked) < limit and any(by_area[area] for area in areas):
        area = areas[index % len(areas)]
        if by_area[area]:
            picked.append(by_area[area].pop())
        index += 1
    return picked


# ----------------------------------------------------------------------- LLM proposal


class _Proposal(BaseModel):
    query: str = Field(description="what the lawyer says to the assistant")
    answer: str = Field(description="the exact value that must appear verbatim in a source doc")
    source_document: str = Field(description="filename of the document that states the answer")


class _Proposals(BaseModel):
    proposals: list[_Proposal]


_SYSTEM = (
    "You write test REQUESTS for a law firm's AI assistant. Given a matter and its source "
    "documents, write what a lawyer would ASK THE ASSISTANT — never a search string.\n\n"
    "This system is used through an assistant, not a search box: the lawyer speaks, and "
    "the assistant decides what to search for. So every request must be something a person "
    "would actually say out loud.\n\n"
    "Write the way a busy lawyer talks — short, natural, conversational. Reference the deal "
    "by the SHORTHAND a lawyer would use: a party name or short deal name ('the Trellis "
    "deal', 'the Meridian ISDA') — NOT a recitation of full legal entity names and document "
    "types.\n"
    "- GOOD: 'What cross-default threshold did we agree for Trellis on the Meridian ISDA?'\n"
    "- GOOD (carries a reference the lawyer holds): 'What's the status on 4478291?'\n"
    "- GOOD: 'Pull up the Hargrove engagement letter — who signed it?'\n"
    "- TOO ROBOTIC: 'In the Trellis Bank AG / Meridian Capital Partners LP matter, what "
    "Cross-Default Threshold does the draft ISDA schedule set for Trellis Bank AG?'\n"
    "- TOO VAGUE (unanswerable — which deal?): 'What is the cross-default threshold?'\n\n"
    "CRITICAL — you are shown ONE matter, but the request will be answered against the "
    "WHOLE FIRM: hundreds of separate matters, most of the same handful of types. There "
    "are dozens of fee letters, purchase agreements, title commitments and NPDES permits "
    "in there, and they are different documents with different values. So a request that "
    "reads unambiguously to you, sitting inside this matter, is unanswerable to someone "
    "who cannot see which matter you meant.\n"
    "EVERY request MUST name its own matter — a party, a deal or project name, a case or "
    "file number. Something that occurs in THIS matter and would not match another.\n"
    "- BROKEN: 'What's the final arrangement fee in the fee letter?' (which of the "
    "dozens of fee letters?)\n"
    "- BROKEN: 'What's the effective date in the purchase agreement?'\n"
    "- BROKEN: anything leaning on 'this deal', 'the agreement', 'our client', 'the "
    "draft' — the reader has no 'this'.\n"
    "- FIXED: 'What's the final arrangement fee in the Westlake fee letter?'\n"
    "- FIXED: 'What's the effective date on the PinnacleAir purchase agreement?'\n"
    "Write the name into the request itself. Do not rely on the document list, the task "
    "title, or anything else you were shown — the lawyer asking has none of it.\n"
    "- FORBIDDEN (this is a search string, not a request): 'hargrove & finch llp', "
    "'$287,614,322.18', 'lf-2024-0917'\n\n"
    "Vary the requests: some should carry a hard anchor the lawyer would have to hand — a "
    "reference number, a case number, a party name — and some should be purely conceptual. "
    "A request may MENTION an anchor, but must never consist only of one.\n\n"
    "The 'answer' MUST appear verbatim (as a substring) in the named source document, and "
    "the request must NOT contain the answer inside it — you are writing the question, not "
    "restating the answer. Never invent values. Prefer requests with a concrete, "
    "distinctive answer (amounts, deadlines, dates, names)."
)


def _build_user_prompt(scenario: dict, doc_texts: dict[str, str], per_scenario: int) -> str:
    listing = []
    for path, text in doc_texts.items():
        name = path.rsplit("/", 1)[-1]
        listing.append(f"### {name}\n{text[:_MAX_DOC_CHARS]}")
    criteria = "\n".join(
        f"- {c.get('title', '')}: {c.get('match_criteria', '')}"
        for c in scenario.get("criteria", [])[:20]
    )
    return (
        f"MATTER: {scenario.get('matter_ref', '')} — {scenario.get('matter_title', '')}\n"
        f"TASK: {scenario.get('title', '')}\n"
        f"INSTRUCTION: {scenario.get('instructions', '')[:1500]}\n\n"
        f"RUBRIC CRITERIA (hint at discriminative values):\n{criteria}\n\n"
        f"SOURCE DOCUMENTS (filename + text excerpt):\n" + "\n\n".join(listing) + "\n\n"
        f"Produce up to {per_scenario} labels (mixed kinds). Put the exact filename from "
        f"the list in 'source_document'. 'answer' must be a verbatim substring of that "
        f"document."
    )


#: (scenario, bundle doc_texts, per_scenario) -> proposals. Injectable for offline tests.
ProposeFn = Callable[[dict, dict[str, str], int], list[_Proposal]]


def _gateway_propose_fn(config: AppConfig, slot: ModelSlot) -> ProposeFn:
    from knowledge_index.pipeline.providers import chat_json

    def propose(scenario: dict, doc_texts: dict[str, str], per_scenario: int) -> list[_Proposal]:
        result = chat_json(
            slot,
            config,
            system=_SYSTEM,
            user=_build_user_prompt(scenario, doc_texts, per_scenario),
            schema=_Proposals,
            # Reasoning models spend part of this budget on hidden reasoning tokens;
            # leave ample room for reasoning + the JSON payload.
            max_output_tokens=16000,
        )
        return result.proposals[:per_scenario]

    return propose


# ------------------------------------------------------------------------ verification


def _longest_prefix_in(answer: str, doc_texts: dict[str, str]) -> tuple[int, str | None]:
    """Largest ``k`` such that ``answer[:k]`` is a substring of some document, and which one.

    A long prefix that matches with a diverging tail means the model reformatted or
    truncated a real value; a near-zero prefix means the answer is genuinely absent.
    """
    best_k, best_doc = 0, None
    for name, text in doc_texts.items():
        low, high = 0, len(answer)
        while low < high:
            mid = (low + high + 1) // 2
            if answer[:mid] in text:
                low = mid
            else:
                high = mid - 1
        if low > best_k:
            best_k, best_doc = low, name
    return best_k, best_doc


def _diagnose_rejection(
    proposal: _Proposal, doc_texts: dict[str, str], unparseable: set[str]
) -> dict:
    """Explain why a proposal failed the verbatim check, for the ``rejected.jsonl`` log."""
    answer = _normalize(proposal.answer)
    if not proposal.query.strip():
        base = {"reason": "empty_query"}
    elif len(answer) < 4:
        base = {"reason": "answer_too_short"}
    else:
        prefix, near_doc = _longest_prefix_in(answer, doc_texts)
        src = proposal.source_document
        src_unreadable = src.rsplit("/", 1)[-1] in unparseable or src.endswith(
            (".xlsx", ".pptx", ".pdf")
        )
        if src_unreadable and prefix < len(answer):
            reason = "answer_in_unreadable_file"  # .xlsx/.pptx/.pdf the grounder can't read
        elif prefix >= 12:
            reason = "reformatted_or_truncated"  # a long prefix matched; the tail diverged
        else:
            reason = "not_in_any_readable_doc"  # essentially absent / composed / hallucinated
        base = {
            "reason": reason,
            "source_hint": src,
            "source_unreadable": src_unreadable,
            "matched_prefix_chars": prefix,
            "answer_chars": len(answer),
            "nearest_doc": near_doc.rsplit("/", 1)[-1] if near_doc else None,
        }
        if prefix >= 8 and near_doc:
            idx = doc_texts[near_doc].find(answer[:prefix])
            base["model_answer_tail"] = answer[prefix : prefix + 40]
            base["doc_actual_text"] = doc_texts[near_doc][idx : idx + prefix + 40]
    return {"query": proposal.query, "answer": proposal.answer, **base}


def _scope_matches(
    answer_normalized: str, principal: str, corpus: dict[str, tuple[str, str]]
) -> list[str]:
    """Every document in the principal's visible scope that states the answer."""
    return sorted(
        path
        for path, (doc_principal, text) in corpus.items()
        if doc_principal == principal and answer_normalized in text
    )


_SPECIFICITY_SYSTEM = (
    "You audit test questions for a document-retrieval benchmark. The corpus is a large "
    "law firm's document store: thousands of documents across hundreds of separate "
    "client matters. Many matters are of the same type, so near-identical documents "
    "recur across them — dozens of fee letters, purchase agreements, title commitments.\n\n"
    "You are shown ONE request a lawyer typed. Decide whether the request, ON ITS OWN, "
    "singles out one specific matter in that corpus.\n"
    "- SPECIFIC: it names a party, deal, project, case number, file reference or other "
    "detail that could only match one matter.\n"
    "- AMBIGUOUS: it leans on context the reader does not have — 'this deal', 'the "
    "agreement', 'the fee letter', 'our client' — and would match many matters equally "
    "well.\n"
    "Judge only what is written. Do not assume a working context, a current matter, or "
    "anything the lawyer might have on screen.\n"
    'Reply as JSON: {"verdict": "SPECIFIC"|"AMBIGUOUS", "reason": "<8 words>"}'
)


def is_self_identifying(
    query: str, config: AppConfig, slot: ModelSlot
) -> tuple[bool, str]:
    """Does this request name the matter it is about, to a reader with no context?

    The generator sees one scenario at a time, so "the draft NPDES" reads as unique to
    it while the corpus holds several. A prompt rule alone does not hold — the previous
    gold carried the same warning and 26% of its requests were still ambiguous, which
    put coin-flip noise straight into every config's score. So the rule is enforced
    here, by a reader that is shown the request and nothing else.

    Fails OPEN: a gateway error must not silently shrink the gold set.
    """
    try:
        message = gateway.complete(
            config,
            slot,
            [
                {"role": "system", "content": _SPECIFICITY_SYSTEM},
                {"role": "user", "content": f"REQUEST: {query}"},
            ],
            max_tokens=2000,
        )
        verdict = gateway.extract_json(message.get("content") or "")
    except Exception:
        return True, "specificity_check_unavailable"
    ok = str(verdict.get("verdict", "")).upper() != "AMBIGUOUS"
    return ok, str(verdict.get("reason", ""))[:80]


def ground_proposal(
    proposal: _Proposal,
    scenario: dict,
    bundle_texts: dict[str, str],
    corpus: dict[str, tuple[str, str]],
    *,
    max_gold_docs: int,
    model: str,
) -> GoldQuery | dict:
    """Verify one proposal; return a ``GoldQuery`` or a rejection record (a dict).

    Three checks, in order:

    1. **The request must not contain its own answer.** A request that restates the
       answer is not a question — and when the gold is "documents containing this
       string", such a label asks a retriever to recompute the label's own
       definition, which a lexical baseline satisfies by construction.
    2. **Verbatim grounding** against the scenario bundle (the only texts the model
       saw), so a hallucinated value cannot become gold.
    3. **Discrimination** across the principal's whole visible scope.

    Gold is then *graded*: the document the answer was drawn from is ``primary``
    (relevance 1.0) and other documents that merely repeat the value are
    ``secondary`` (0.3). Flat labelling made a system that returned THE answering
    document at rank 1 score ~0.64 instead of 1.0, and rewarded retrieving
    incidental co-mentions.
    """
    answer = _normalize(proposal.answer)
    unparseable = {
        path.rsplit("/", 1)[-1]
        for path in scenario["document_paths"]
        if path not in bundle_texts
    }
    if len(answer) < 4 or not proposal.query.strip():
        return _diagnose_rejection(proposal, bundle_texts, unparseable)

    query_normalized = _normalize(proposal.query)
    if answer in query_normalized:
        return {
            "query": proposal.query,
            "answer": proposal.answer,
            "reason": "request_restates_its_own_answer",
        }

    if not any(answer in text for text in bundle_texts.values()):
        return _diagnose_rejection(proposal, bundle_texts, unparseable)

    matches = _scope_matches(answer, scenario["principal"], corpus)
    if len(matches) > max_gold_docs:
        return {
            "query": proposal.query,
            "answer": proposal.answer,
            "reason": "not_discriminative",
            "scope_matches": len(matches),
            "max_gold_docs": max_gold_docs,
        }

    # The model names the document it drew the answer from; trust it only if that
    # document actually contains the answer, else fall back to the first match.
    hint = proposal.source_document.rsplit("/", 1)[-1].casefold()
    primary = next(
        (path for path in matches if path.rsplit("/", 1)[-1].casefold() == hint), matches[0]
    )
    return GoldQuery(
        id="",  # assigned by the caller once accepted
        kind="request",
        query=proposal.query.strip(),
        gold_paths=matches,
        principals=[scenario["principal"]],
        matter_ref=scenario["matter_ref"],
        practice_area=scenario["practice_area"],
        meta={
            "answer": proposal.answer,
            "anchor": classify_anchor(proposal.query, proposal.answer),
            "primary_path": primary,
            "secondary_paths": [path for path in matches if path != primary],
            "source_hint": proposal.source_document,
            "scope_matches": len(matches),
            "extracted_by": f"{model}/{PROMPT_VERSION}",
        },
    )


# --------------------------------------------------------------------------- pipeline


def generate_gold(
    corpus_dir: str | Path,
    config: AppConfig | None = None,
    *,
    per_scenario: int = 4,
    model_slot: ModelSlot | None = None,
    limit_scenarios: int | None = None,
    seed: int = 42,
    max_gold_docs: int = 5,
    concurrency: int = 64,
    output_dir: str | Path | None = None,
    propose_fn: ProposeFn | None = None,
) -> dict:
    """Generate, verify, and append gold to ``retrieval-gold.jsonl``.

    Labels are deduplicated by ``(kind, normalized query)`` against the whole file,
    so this is safe to re-run and to grow incrementally (e.g. a higher
    ``limit_scenarios`` on a second pass). Rejections are appended to
    ``rejected.jsonl`` with diagnosed reasons.

    Proposal calls fan out over a bounded thread pool (``concurrency`` — vLLM/gateway
    backends batch concurrent requests); verification runs serially in scenario
    order, so the output and dedupe behavior are deterministic regardless of the
    pool size.

    Each scenario is proposed, verified and **written** as an independent unit, so a
    kill or a crash costs at most the scenarios in flight, and re-running skips every
    scenario already present in the gold file.

    ``output_dir`` is where ``retrieval-gold.jsonl`` / ``rejected.jsonl`` are written
    (and read for dedupe). Defaults to ``corpus_dir``; pass a writable location when
    the corpus mount is read-only (e.g. ``/testdata`` in the app container).
    """
    corpus_dir = Path(corpus_dir).resolve()
    if propose_fn is None:
        if config is None:
            raise ValueError("config is required unless a propose_fn is injected")
        slot = model_slot or config.models.judge
        propose_fn = _gateway_propose_fn(config, slot)
        model_name = slot.model
    else:
        model_name = "injected"
        slot = model_slot  # may be None: an injected propose_fn skips the LLM gate

    scenarios = load_scenarios(corpus_dir)
    corpus = corpus_texts(corpus_dir, scenarios)
    sampled = stratified_sample(scenarios, limit_scenarios, seed=seed)

    out_dir = Path(output_dir).resolve() if output_dir else corpus_dir
    out_dir.mkdir(parents=True, exist_ok=True)
    gold_path = out_dir / "retrieval-gold.jsonl"
    existing = (
        [json.loads(line) for line in gold_path.read_text(encoding="utf-8").splitlines() if line]
        if gold_path.exists()
        else []
    )
    seen = {_normalize(item["query"]) for item in existing}
    # Scenario ids already on disk: their work is done and must not be redone.
    done_scenarios = {item["id"].split("#", 1)[0] for item in existing if "#" in item["id"]}

    accepted: list[GoldQuery] = []
    rejected: list[dict] = []
    stats = {
        "scenarios": len(sampled),
        "corpus_documents": len(corpus),
        "proposed": 0,
        "accepted": 0,
        "duplicate": 0,
        "errors": 0,
    }
    worklist = [
        (scenario, bundle_texts)
        for scenario in sampled
        if (
            bundle_texts := {
                path: corpus[path][1] for path in scenario["document_paths"] if path in corpus
            }
        )
        # resume: a scenario already represented in the gold file is finished work
        and scenario["scenario_id"] not in done_scenarios
    ]
    stats["scenarios"] = len(worklist)
    stats["resumed_scenarios"] = len(done_scenarios)

    write_lock = threading.Lock()
    gold_handle = gold_path.open("a", encoding="utf-8")
    rejected_handle = (out_dir / "rejected.jsonl").open("a", encoding="utf-8")

    def _process(entry: tuple[dict, dict[str, str]]) -> None:
        """Propose, verify and PERSIST one scenario. Independent unit of work: a
        crash or a kill costs at most the scenarios in flight, and re-running skips
        everything already on disk."""
        scenario, bundle_texts = entry
        try:
            proposals = propose_fn(scenario, bundle_texts, per_scenario)
        except Exception:  # gateway/schema failure on one scenario must not abort
            with write_lock:
                stats["errors"] += 1
            return
        for index, proposal in enumerate(proposals[:per_scenario]):
            grounded = ground_proposal(
                proposal,
                scenario,
                bundle_texts,
                corpus,
                max_gold_docs=max_gold_docs,
                model=model_name,
            )
            # Only worth an LLM call once the cheap checks have passed.
            if not isinstance(grounded, dict) and config is not None and slot is not None:
                specific, why = is_self_identifying(grounded.query, config, slot)
                if not specific:
                    grounded = {
                        "query": grounded.query,
                        "answer": grounded.meta.get("answer"),
                        "reason": "not_self_identifying",
                        "detail": why,
                    }
            with write_lock:
                stats["proposed"] += 1
                if isinstance(grounded, dict):
                    grounded["scenario_id"] = scenario["scenario_id"]
                    grounded["matter_ref"] = scenario["matter_ref"]
                    rejected.append(grounded)
                    rejected_handle.write(json.dumps(grounded, ensure_ascii=False) + "\n")
                    continue
                key = _normalize(grounded.query)
                if key in seen:
                    stats["duplicate"] += 1
                    continue
                seen.add(key)
                grounded.id = f"{scenario['scenario_id']}#{index + 1}"
                accepted.append(grounded)
                stats["accepted"] += 1
                gold_handle.write(json.dumps(asdict(grounded), ensure_ascii=False) + "\n")
                gold_handle.flush()

    try:
        with ThreadPoolExecutor(max_workers=max(1, concurrency)) as pool:
            list(pool.map(_process, worklist))
    finally:
        gold_handle.close()
        rejected_handle.close()

    stats["gold_path"] = str(gold_path)
    stats["rejected"] = len(rejected)
    stats["reject_reasons"] = dict(Counter(record["reason"] for record in rejected))
    stats["by_anchor"] = dict(Counter(item.meta["anchor"] for item in accepted))
    return stats
