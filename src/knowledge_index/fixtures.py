"""Deterministic, fictional German law-firm DMS fixture generator."""

from __future__ import annotations

import hashlib
import json
import random
import zipfile
from dataclasses import dataclass, field
from datetime import UTC, datetime
from email.message import EmailMessage
from pathlib import Path

from docx import Document as WordDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


@dataclass
class FixtureRecord:
    relative_path: str
    matter_ref: str | None
    doc_type: str
    logical_document: str
    version: dict | None = None
    relations: list[dict] = field(default_factory=list)
    acl: list[dict] = field(default_factory=list)
    expected_pipeline: str = "done"
    expected_error: str | None = None
    rationale: dict | None = None
    pii: list[str] = field(default_factory=list)
    content_hash: str | None = None
    # Ontology-tolerant acceptance: entries are ancestor LABELS of the active
    # ontology ("Agreements", "Litigation Document") — the verifier passes when
    # any ancestor of the typed node carries one of them. "*" (or an empty
    # list) only requires the document to be typed at some visible node; raw
    # LMSS has no home for several everyday genres (emails, notes, judgments).
    accepted_doc_types: list[str] = field(default_factory=list)


def generate_mock_dms(output: str | Path, *, seed: int = 42) -> dict:
    output = Path(output).resolve()
    if output.exists() and any(output.iterdir()):
        raise ValueError(f"fixture output must be empty: {output}")
    source_root = output / "mock_dms"
    source_root.mkdir(parents=True, exist_ok=True)
    random.Random(seed)  # reserved for deterministic scale knobs

    ma_acl = [_grant("group:ma-team")]
    litigation_acl = [_grant("group:litigation")]
    internal_acl = [_grant("group:all-lawyers")]
    records: list[FixtureRecord] = []

    ma = source_root / "Mandate" / "M-2026-0042 Projekt Falke"
    ma.mkdir(parents=True)
    draft = ma / "Unternehmenskaufvertrag_Entwurf_v1.docx"
    _contract_docx(
        draft,
        title="Entwurf Unternehmenskaufvertrag — Projekt Falke",
        liability="Die Haftung der Verkäuferin ist unbeschränkt.",
    )
    records.append(
        FixtureRecord(
            relative_path=_relative(draft, source_root),
            matter_ref="M-2026-0042",
            doc_type="share_purchase_agreement",
            accepted_doc_types=["Agreements", "Transactional Document"],
            logical_document="falke-spa",
            version={"ordinal": 1, "status": "draft"},
            acl=ma_acl,
        )
    )

    redline = ma / "Unternehmenskaufvertrag_redline_v2.docx"
    _redline_contract(redline)
    records.append(
        FixtureRecord(
            relative_path=_relative(redline, source_root),
            matter_ref="M-2026-0042",
            doc_type="share_purchase_agreement",
            accepted_doc_types=["Agreements", "Transactional Document"],
            logical_document="falke-spa",
            version={"ordinal": 2, "status": "draft", "previous": draft.name},
            acl=ma_acl,
            rationale={
                "locus_keywords": ["§ 9", "haftung"],
                "accepted_categories": [
                    "legal_risk",
                    "negotiation_concession",
                    "market_standard",
                ],
                "required_keywords": ["haftung", "kaufpreis"],
                "expect_record": True,
            },
            pii=["Falke Erwerbs GmbH", "Adler Industrie GmbH", "Dr. Anna Beispiel"],
        )
    )

    final = ma / "Unternehmenskaufvertrag_final.docx"
    _contract_docx(
        final,
        title="Unternehmenskaufvertrag — Projekt Falke",
        liability="Die Haftung der Verkäuferin ist auf den Kaufpreis begrenzt.",
    )
    records.append(
        FixtureRecord(
            relative_path=_relative(final, source_root),
            matter_ref="M-2026-0042",
            doc_type="share_purchase_agreement",
            accepted_doc_types=["Agreements", "Transactional Document"],
            logical_document="falke-spa",
            version={"ordinal": 3, "status": "final", "previous": redline.name},
            acl=ma_acl,
        )
    )

    executed = ma / "Unternehmenskaufvertrag_unterzeichnet.txt"
    executed.write_text(
        "Unternehmenskaufvertrag — unterzeichnete Fassung\n"
        "M-2026-0042\n"
        "zwischen der Falke Erwerbs GmbH (Käuferin) und der Adler Industrie GmbH "
        "(Verkäuferin)\n"
        "§ 9 Haftung\nDie Haftung der Verkäuferin ist auf den Kaufpreis begrenzt.\n"
        "Unterzeichnet am 17. Mai 2026.",
        encoding="utf-8",
    )
    records.append(
        FixtureRecord(
            relative_path=_relative(executed, source_root),
            matter_ref="M-2026-0042",
            doc_type="share_purchase_agreement",
            accepted_doc_types=["Agreements", "Transactional Document"],
            logical_document="falke-spa",
            version={"ordinal": 4, "status": "executed", "previous": final.name},
            acl=ma_acl,
            pii=["Falke Erwerbs GmbH", "Adler Industrie GmbH"],
        )
    )

    annex = ma / "Anlage_1_Garantiekatalog.txt"
    annex.write_text(
        "M-2026-0042\nAnlage 1 zum Unternehmenskaufvertrag\nGarantien der Verkäuferin.\n",
        encoding="utf-8",
    )
    records.append(
        FixtureRecord(
            relative_path=_relative(annex, source_root),
            matter_ref="M-2026-0042",
            doc_type="other_annex",
            accepted_doc_types=["*"],
            logical_document="falke-annex-1",
            relations=[{"kind": "annex_of", "target": "falke-spa"}],
            acl=ma_acl,
        )
    )

    mail_dir = ma / "Korrespondenz"
    mail_dir.mkdir()
    first_mail = mail_dir / "2026-05-10_Haftung.eml"
    _email(
        first_mail,
        subject="M-2026-0042 – Haftungsbegrenzung",
        message_id="<falke-1@mock.kanzlei>",
        body=(
            "Die unbeschränkte Haftung ist aus rechtlicher Risikosicht nicht vertretbar. "
            "Bitte auf den Kaufpreis begrenzen; dies ist keine bloße Mandantenpräferenz."
        ),
    )
    records.append(
        FixtureRecord(
            relative_path=_relative(first_mail, source_root),
            matter_ref="M-2026-0042",
            doc_type="email",
            logical_document="falke-mail-1",
            acl=ma_acl,
            rationale={
                "locus": "§ 9 Haftung",
                "category": "legal_risk",
                "generalizable": True,
            },
        )
    )

    reply = mail_dir / "2026-05-11_AW_Haftung.eml"
    _email(
        reply,
        subject="AW: M-2026-0042 – Haftungsbegrenzung",
        message_id="<falke-2@mock.kanzlei>",
        in_reply_to="<falke-1@mock.kanzlei>",
        body="Einverstanden. Für die Zahlungsfrist wünscht die Mandantin dagegen zehn Tage.",
    )
    records.append(
        FixtureRecord(
            relative_path=_relative(reply, source_root),
            matter_ref="M-2026-0042",
            doc_type="email",
            logical_document="falke-mail-2",
            relations=[{"kind": "responds_to", "target": "falke-mail-1"}],
            acl=ma_acl,
            rationale={"category": "client_instruction", "generalizable": False},
        )
    )

    litigation = source_root / "Mandate" / "M-2026-0099 Müller ._ Schmidt"
    litigation.mkdir(parents=True)
    pleading = litigation / "Klageschrift_final.docx"
    _simple_docx(
        pleading,
        "Klageschrift M-2026-0099",
        [
            "Müller Handels GmbH gegen Schmidt Logistik AG",
            "Die Beklagte wird verurteilt, die offene Kaufpreisforderung zu zahlen.",
            "Beweis: Anlage K1 und Urteil BGH VIII ZR 1/25.",
        ],
    )
    records.append(
        FixtureRecord(
            relative_path=_relative(pleading, source_root),
            matter_ref="M-2026-0099",
            doc_type="statement_of_claim",
            logical_document="mueller-complaint",
            version={"ordinal": 1, "status": "final"},
            relations=[{"kind": "references", "target": "mueller-judgment"}],
            acl=litigation_acl,
            pii=["Müller Handels GmbH", "Schmidt Logistik AG"],
        )
    )
    evidence = litigation / "Anlage_K1_Kaufvertrag.txt"
    evidence.write_text(
        "M-2026-0099\nKaufvertrag über Logistikleistungen vom 4. Februar 2026.\n",
        encoding="utf-8",
    )
    records.append(
        FixtureRecord(
            relative_path=_relative(evidence, source_root),
            matter_ref="M-2026-0099",
            doc_type="other_annex",
            accepted_doc_types=["*"],
            logical_document="mueller-k1",
            relations=[{"kind": "annex_of", "target": "mueller-complaint"}],
            acl=litigation_acl,
        )
    )

    response = litigation / "Klageerwiderung_final.docx"
    _simple_docx(
        response,
        "Klageerwiderung M-2026-0099",
        [
            "Schmidt Logistik AG gegen Müller Handels GmbH",
            "Die Beklagte beantragt, die Klage abzuweisen.",
            "Die behauptete Kaufpreisforderung ist wegen Mängeln nicht fällig.",
        ],
    )
    records.append(
        FixtureRecord(
            relative_path=_relative(response, source_root),
            matter_ref="M-2026-0099",
            doc_type="statement_of_defense",
            logical_document="mueller-response",
            version={"ordinal": 1, "status": "final"},
            relations=[{"kind": "responds_to", "target": "mueller-complaint"}],
            acl=litigation_acl,
            pii=["Müller Handels GmbH", "Schmidt Logistik AG"],
        )
    )

    judgment = litigation / "Urteil_BGH_VIII_ZR_1-25_final.txt"
    judgment.write_text(
        "Urteil BGH VIII ZR 1/25\nM-2026-0099\n"
        "Bei behebbaren Mängeln kann die Fälligkeit der Kaufpreisforderung ausgesetzt sein.\n",
        encoding="utf-8",
    )
    records.append(
        FixtureRecord(
            relative_path=_relative(judgment, source_root),
            matter_ref="M-2026-0099",
            doc_type="judgment",
            logical_document="mueller-judgment",
            version={"ordinal": 1, "status": "final"},
            acl=litigation_acl,
        )
    )

    litigation_mail = litigation / "2026-06-02_Maengelanzeige.eml"
    _email(
        litigation_mail,
        subject="M-2026-0099 – Mängelanzeige und Kaufpreisforderung",
        message_id="<mueller-1@mock.kanzlei>",
        body=(
            "Die Mandantin bestreitet die Fälligkeit wegen der dokumentierten Mängel. "
            "Bitte die Klageerwiderung mit dem Urteil BGH VIII ZR 1/25 abstimmen."
        ),
    )
    records.append(
        FixtureRecord(
            relative_path=_relative(litigation_mail, source_root),
            matter_ref="M-2026-0099",
            doc_type="email",
            logical_document="mueller-mail-1",
            acl=litigation_acl,
        )
    )

    internal = source_root / "Kanzlei Intern"
    internal.mkdir()
    policy = internal / "KI_Richtlinie_final.txt"
    policy.write_text(
        "Kanzleiinterne Richtlinie: Mandatsdaten dürfen nur innerhalb der jeweiligen "
        "Berechtigungsgruppe verarbeitet werden.",
        encoding="utf-8",
    )
    records.append(
        FixtureRecord(
            relative_path=_relative(policy, source_root),
            matter_ref=None,
            doc_type="note",
            accepted_doc_types=["*"],
            logical_document="internal-ai-policy",
            acl=internal_acl,
        )
    )

    duplicate = internal / "Falsch_abgelegt_Vertrag_final.docx"
    duplicate.write_bytes(final.read_bytes())
    records.append(
        FixtureRecord(
            relative_path=_relative(duplicate, source_root),
            matter_ref="M-2026-0042",
            doc_type="share_purchase_agreement",
            accepted_doc_types=["Agreements", "Transactional Document"],
            logical_document="falke-spa",
            version={"ordinal": 3, "status": "final", "duplicate_of": final.name},
            acl=internal_acl,
        )
    )

    poison = source_root / "Eingang unsortiert" / "M-2026-0099_scan_alt.bin"
    poison.parent.mkdir()
    poison.write_bytes(b"\x00\xffnot-a-supported-document\x00")
    records.append(
        FixtureRecord(
            relative_path=_relative(poison, source_root),
            matter_ref="M-2026-0099",
            doc_type="unknown",
            logical_document="poison-legacy-scan",
            acl=litigation_acl,
            expected_pipeline="quarantined",
            expected_error="UnsupportedDocument",
        )
    )

    for record in records:
        _freeze_zip_timestamps(source_root / record.relative_path)
        record.content_hash = _sha256(source_root / record.relative_path)

    manifest_path = output / "ground-truth.jsonl"
    manifest_path.write_text(
        "".join(json.dumps(record.__dict__, ensure_ascii=False) + "\n" for record in records),
        encoding="utf-8",
    )
    acl_path = output / "acl-by-path.json"
    acl_path.write_text(
        json.dumps({record.relative_path: record.acl for record in records}, indent=2),
        encoding="utf-8",
    )
    scenario = {
        "seed": seed,
        "source_root": str(source_root),
        "manifest": str(manifest_path),
        "acl_by_path": str(acl_path),
        "file_count": len(records),
        "matter_refs": ["M-2026-0042", "M-2026-0099"],
    }
    (output / "scenario.json").write_text(
        json.dumps(scenario, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return scenario


def _grant(principal: str) -> dict:
    return {"principal": principal, "principal_kind": "group", "access": "allow"}


def _relative(path: Path, root: Path) -> str:
    return path.relative_to(root).as_posix()


# OOXML files are zip archives, and python-docx stamps each entry with the wall-clock
# time at save, which would make two runs of the same seed differ. The timestamp carries
# no fixture meaning, so it is pinned to the same fictional date as the content.
FIXTURE_ZIP_TIMESTAMP = (2026, 5, 10, 10, 0, 0)


def _freeze_zip_timestamps(path: Path) -> None:
    """Rewrite an OOXML file with fixed entry timestamps. No-op for anything else."""
    if not zipfile.is_zipfile(path):
        return
    with zipfile.ZipFile(path) as archive:
        entries = [(info, archive.read(info.filename)) for info in archive.infolist()]
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
        for info, payload in entries:
            frozen = zipfile.ZipInfo(info.filename, date_time=FIXTURE_ZIP_TIMESTAMP)
            frozen.compress_type = info.compress_type
            frozen.external_attr = info.external_attr
            archive.writestr(frozen, payload)


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _simple_docx(path: Path, title: str, paragraphs: list[str]) -> None:
    document = WordDocument()
    document.core_properties.title = title
    document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)


def _contract_docx(path: Path, *, title: str, liability: str) -> None:
    _simple_docx(
        path,
        title,
        [
            "M-2026-0042",
            "zwischen der Falke Erwerbs GmbH (Käuferin) und der Adler Industrie GmbH (Verkäuferin)",
            "§ 1 Kaufgegenstand",
            "Die Verkäuferin verkauft sämtliche Geschäftsanteile an der Zielgesellschaft.",
            "§ 9 Haftung",
            liability,
            "§ 14 Schlussbestimmungen",
            "Es gilt deutsches Recht.",
        ],
    )


def _redline_contract(path: Path) -> None:
    document = WordDocument()
    document.core_properties.title = "Unternehmenskaufvertrag Redline"
    document.add_heading("Unternehmenskaufvertrag — Redline v2", level=1)
    document.add_paragraph("M-2026-0042\n§ 9 Haftung")
    paragraph = document.add_paragraph("Die Haftung der Verkäuferin ist ")
    deleted = OxmlElement("w:del")
    deleted.set(qn("w:id"), "1")
    deleted.set(qn("w:author"), "Dr. Anna Beispiel")
    deleted.set(qn("w:date"), "2026-05-10T10:00:00Z")
    deleted_run = OxmlElement("w:r")
    deleted_text = OxmlElement("w:delText")
    deleted_text.text = "unbeschränkt"
    deleted_run.append(deleted_text)
    deleted.append(deleted_run)
    inserted = OxmlElement("w:ins")
    inserted.set(qn("w:id"), "2")
    inserted.set(qn("w:author"), "Dr. Anna Beispiel")
    inserted.set(qn("w:date"), "2026-05-10T10:01:00Z")
    inserted_run = OxmlElement("w:r")
    inserted_text = OxmlElement("w:t")
    inserted_text.text = "auf den Kaufpreis begrenzt"
    inserted_run.append(inserted_text)
    inserted.append(inserted_run)
    paragraph._p.append(deleted)
    paragraph._p.append(inserted)
    paragraph.add_run(".")
    document.save(path)


def _email(
    path: Path,
    *,
    subject: str,
    message_id: str,
    body: str,
    in_reply_to: str | None = None,
) -> None:
    message = EmailMessage()
    message["From"] = "anna.beispiel@kanzlei.invalid"
    message["To"] = "ben.fiktiv@kanzlei.invalid"
    message["Subject"] = subject
    message["Message-ID"] = message_id
    message["Date"] = datetime(2026, 5, 10, 10, 0, tzinfo=UTC)
    if in_reply_to:
        message["In-Reply-To"] = in_reply_to
    message.set_content(body)
    path.write_bytes(message.as_bytes())
