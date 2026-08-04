"""Benchmark harness: offline unit coverage + one live end-to-end run.

The unit tests need no database, models, or docker stack — the metric core takes an
injectable search function and the corpus/gold builders are pure filesystem work.
The single ``@pytest.mark.integration`` test packs a tiny corpus, ingests it through
the real pipeline, and asserts the full system beats the naive-dense baseline.
"""

from __future__ import annotations

import json
from dataclasses import dataclass
from email.message import EmailMessage
from pathlib import Path

import pytest
from docx import Document as WordDocument

from knowledge_index.benchmark import metrics, presets
from knowledge_index.benchmark.gold import generate_gold
from knowledge_index.benchmark.task_corpus import build_task_corpus
from knowledge_index.benchmark.retrieval_eval import run_queries
from knowledge_index.config import AppConfig
from tests.conftest import TEST_EMBEDDING_MODEL, TEST_LLM_MODEL


# --------------------------------------------------------------------------- metrics


def test_metrics_match_hand_computed_values() -> None:
    gold = {"a", "b"}
    ranked_covers = [{"a"}, set(), {"b"}]  # rank1 a, rank2 miss, rank3 b
    assert metrics.recall_at_k(ranked_covers, gold, 1) == 0.5
    assert metrics.recall_at_k(ranked_covers, gold, 5) == 1.0
    assert metrics.precision_at_k(ranked_covers, gold, 1) == 1.0
    assert metrics.reciprocal_rank(ranked_covers, gold) == 1.0
    # dcg = 1/log2(2) + 1/log2(4) = 1.5 ; ideal = 1/log2(2) + 1/log2(3)
    assert metrics.ndcg_at_k(ranked_covers, gold, 10) == pytest.approx(0.9197, abs=1e-3)


def test_metrics_reward_earlier_hits() -> None:
    gold = {"a"}
    early = metrics.ndcg_at_k([{"a"}, set(), set()], gold, 10)
    late = metrics.ndcg_at_k([set(), set(), {"a"}], gold, 10)
    assert early > late == pytest.approx(0.5, abs=1e-3)


def test_graded_gold_splits_primary_from_co_mentions() -> None:
    item = {"gold_paths": ["a", "b", "c"], "meta": {"primary_path": "b"}}
    primary, gains = metrics.graded_gold(item)
    assert primary == {"b"}
    assert gains == {"a": 0.3, "b": 1.0, "c": 0.3}

    # Gold written before grading existed degrades to flat relevance, not to zero.
    assert metrics.graded_gold({"gold_paths": ["a", "b"], "meta": {}}) == (
        {"a", "b"},
        {"a": 1.0, "b": 1.0},
    )
    # A primary_path that no longer appears in gold_paths must not empty the set.
    assert metrics.graded_gold({"gold_paths": ["a"], "meta": {"primary_path": "gone"}})[0] == {"a"}


def test_graded_scoring_ranks_the_answering_document_above_co_mentions() -> None:
    item = {"gold_paths": ["a", "b", "c"], "meta": {"primary_path": "b"}}
    primary, gains = metrics.graded_gold(item)
    gold = set(item["gold_paths"])

    # A co-mention at rank 1 is worth something, but it is not a hit.
    weak = metrics.QueryScore.compute("q", "none", [{"c"}, {"b"}], gold, primary=primary, gains=gains)
    assert weak.recall[1] == 0.0
    assert weak.ndcg[1] == pytest.approx(0.3, abs=1e-3)
    assert weak.mrr == 0.5

    # Answering document first, co-mentions behind it, is a perfect ranking.
    best = metrics.QueryScore.compute(
        "q", "none", [{"b"}, {"a"}, {"c"}], gold, primary=primary, gains=gains
    )
    assert best.recall[1] == 1.0
    assert best.ndcg[10] == pytest.approx(1.0, abs=1e-6)
    assert best.gold_size == 3 and best.primary_size == 1


def test_aggregate_splits_by_kind() -> None:
    scores = [
        metrics.QueryScore.compute("q1", "factoid", [{"a"}], {"a"}),
        metrics.QueryScore.compute("q2", "instruction_working_set", [set(), {"b"}], {"b"}),
    ]
    report = metrics.aggregate(scores)
    assert report["overall"]["queries"] == 2
    assert set(report["by_kind"]) == {"factoid", "instruction_working_set"}
    assert report["by_kind"]["factoid"]["recall"]["@1"] == 1.0


def test_paired_bootstrap_detects_a_real_lift_and_rejects_noise() -> None:
    # a constant +0.2 lift over 40 queries is unambiguous
    reference = [0.5] * 40
    lifted = [0.7] * 40
    result = metrics.paired_bootstrap(lifted, reference)
    assert result["delta"] == pytest.approx(0.2)
    assert result["significant"] is True
    assert result["ci95"][0] > 0

    # symmetric noise around zero must not read as significant
    noise = [0.5 + (0.1 if i % 2 else -0.1) for i in range(40)]
    result = metrics.paired_bootstrap(noise, reference)
    assert result["significant"] is False

    # deterministic across calls (seeded)
    assert metrics.paired_bootstrap(lifted, reference) == metrics.paired_bootstrap(
        lifted, reference
    )

    with pytest.raises(ValueError):
        metrics.paired_bootstrap([1.0], [1.0, 2.0])
    assert metrics.paired_bootstrap([], [])["n"] == 0


# --------------------------------------------------------------------------- presets


def test_naive_dense_zeroes_the_smart_legs_without_mutating_input() -> None:
    config = AppConfig()
    before = config.retrieval.model_dump()
    ablated = presets.apply_preset(config, "naive_dense")
    assert ablated.retrieval.weight_lexical == 0.0
    assert ablated.retrieval.weight_identifier == 0.0
    assert ablated.retrieval.weight_semantic == 1.0
    assert ablated.retrieval.collapse_per_document is False
    assert ablated.retrieval.rerank_enabled is False
    # a generic stack has no profile/clause rows — body chunks only
    assert ablated.retrieval.search_chunk_kinds == ["chunk"]
    # the original config is untouched (deep copy)
    assert config.retrieval.model_dump() == before
    assert config.retrieval.weight_identifier == 1.5


def test_full_preset_is_the_shipped_defaults() -> None:
    config = AppConfig()
    assert presets.apply_preset(config, "full").retrieval.model_dump() == (
        config.retrieval.model_dump()
    )


def test_ablations_revert_exactly_one_feature() -> None:
    config = AppConfig()
    no_identifier = presets.apply_preset(config, "full_no_identifier")
    assert no_identifier.retrieval.weight_identifier == 0.0
    assert no_identifier.retrieval.collapse_per_document is True  # everything else intact

    no_profiles = presets.apply_preset(config, "full_no_profiles")
    assert no_profiles.retrieval.search_chunk_kinds == ["chunk", "clause"]
    assert no_profiles.retrieval.weight_identifier == 1.5

    rerank = presets.apply_preset(config, "full_rerank")
    assert rerank.retrieval.rerank_enabled is True


def test_preset_groups_resolve_and_always_include_full() -> None:
    assert "full" in presets.resolve_presets("competitors")
    assert presets.resolve_presets("bm25") == ("bm25", "full")
    assert presets.resolve_presets("full,bm25") == ("full", "bm25")
    assert set(presets.resolve_presets("all")) == set(presets.PRESETS)
    with pytest.raises(KeyError):
        presets.resolve_presets("typo")
    with pytest.raises(KeyError):
        presets.apply_preset(AppConfig(), "typo")


# ---------------------------------------------------------------- corpus + gold build


def _write_task(task_dir: Path, *, title: str, marker: str, filler: str) -> None:
    """One upstream-shaped task: task.json + a docx (holds the marker) + an email."""
    scenario = task_dir / "scenario-01"
    documents = scenario / "documents"
    documents.mkdir(parents=True)
    document = WordDocument()
    document.add_paragraph(f"{title} playbook")
    document.add_paragraph(f"The Secured Party is {marker} acting as Administrative Agent.")
    document.save(documents / "playbook.docx")
    email = EmailMessage()
    email["Subject"] = f"{title} engagement"
    email["From"] = "partner@firm.invalid"
    email["To"] = "associate@firm.invalid"
    email.set_content(f"Please prepare the {title}. {filler}")
    (documents / "engagement.eml").write_bytes(email.as_bytes())
    (scenario / "task.json").write_text(
        json.dumps(
            {
                "title": f"Draft {title}",
                "instructions": (
                    f"Using the provided template, draft the {title} using the reference "
                    "files.\n\n### Output:\nresult.docx"
                ),
                "criteria": [
                    {
                        "id": "C-001",
                        "title": f"{marker} identified in preamble",
                        "match_criteria": (
                            f"PASS if the preamble identifies the Secured Party as "
                            f"'{marker}'. FAIL otherwise."
                        ),
                    }
                ],
            }
        ),
        encoding="utf-8",
    )


def _mini_task_set(root: Path) -> Path:
    area = root / "tasks" / "contracts" / "banking"
    _write_task(
        area / "account-control-agreement-first-draft",
        title="Account Control Agreement",
        marker="Pinnacle Capital Partners, LLC",
        filler="It is a Delaware limited liability company.",
    )
    _write_task(
        area / "credit-support-annex-first-draft",
        title="Credit Support Annex",
        marker="Meridian Engineered Systems, Inc.",
        filler="It is a Virginia corporation.",
    )
    return root


def _write_flat_task(task_dir: Path, *, title: str) -> None:
    """Write the flat task/task.json layout used by current upstream areas."""
    documents = task_dir / "documents"
    documents.mkdir(parents=True)
    document = WordDocument()
    document.add_heading(title, 0)
    document.add_paragraph(f"Reference material for {title}.")
    document.save(documents / "reference.docx")
    (task_dir / "task.json").write_text(
        json.dumps({"title": title, "instructions": f"Review {title}.", "criteria": []}),
        encoding="utf-8",
    )


def test_build_corpus_shapes_a_firm_tree_and_is_deterministic(tmp_path: Path) -> None:
    source = _mini_task_set(tmp_path / "task_set")
    summary = build_task_corpus(tmp_path / "out", source, areas=["contracts/banking"])

    assert summary["matters"] == 2
    assert summary["scenarios"] == 2
    assert summary["documents"] == 4
    assert summary["principals"] == ["group:contracts-banking-team"]
    # the connector-facing ACL map covers every packed document, under Mandate/
    acl = json.loads(Path(summary["acl_by_path"]).read_text(encoding="utf-8"))
    assert len(acl) == 4
    assert all(path.startswith("Mandate/") for path in acl)

    repeat = build_task_corpus(tmp_path / "out2", source, areas=["contracts/banking"])
    assert repeat["content_hash"] == summary["content_hash"]


def test_build_corpus_supports_flat_task_layout(tmp_path: Path) -> None:
    source = tmp_path / "task_set"
    area = source / "tasks" / "banking-finance"
    _write_flat_task(area / "draft-fee-letter", title="Fee Letter")
    _write_flat_task(area / "review-credit-agreement", title="Credit Agreement")

    summary = build_task_corpus(tmp_path / "out", source, areas=["banking-finance"])

    assert summary["matters"] == 2
    assert summary["scenarios"] == 2
    assert summary["documents"] == 2
    assert summary["principals"] == ["group:banking-finance-team"]


def test_firm_layout_uses_the_injected_party_resolver(tmp_path: Path) -> None:
    """The firm builder names each matter via the injected resolver and groups matters
    under one client folder + one per-client ACL — the seam the LLM resolver plugs into."""
    source = _mini_task_set(tmp_path / "task_set")
    seen: list[int] = []

    def resolver(scenarios):  # both matters belong to one client, distinct counterparties
        seen.append(len(scenarios))
        return [("Northwind", "Acme"), ("Northwind", "Globex")]

    summary = build_task_corpus(
        tmp_path / "out",
        source,
        areas=["contracts/banking"],
        layout="firm",
        resolve_parties=resolver,
    )

    assert seen == [2]  # the resolver saw exactly the selected scenarios
    assert summary["layout"] == "firm"
    assert summary["clients"] == 1
    assert summary["client_names"] == ["Northwind"]
    assert summary["matters"] == 2
    assert summary["principals"] == ["group:northwind"]
    acl = json.loads(Path(summary["acl_by_path"]).read_text(encoding="utf-8"))
    assert acl and all(path.startswith("Clients/Northwind/") for path in acl)
    # each matter is "<instrument> — <counterparty>"; pairing follows the seeded order
    titles = {r["matter_title"] for r in _read_scenarios(summary)}
    assert {t.split(" — ")[0] for t in titles} == {
        "Account Control Agreement",
        "Credit Support Annex",
    }
    assert {t.split(" — ")[1] for t in titles} == {"Acme", "Globex"}


def test_firm_naming_resolver_rejected_for_flat_layout(tmp_path: Path) -> None:
    source = _mini_task_set(tmp_path / "task_set")
    with pytest.raises(ValueError, match="firm layout"):
        build_task_corpus(
            tmp_path / "out",
            source,
            areas=["contracts/banking"],
            layout="flat",
            resolve_parties=lambda scenarios: [("X", "Y")],
        )


def test_firm_layout_applies_a_curated_structure_manifest(tmp_path: Path) -> None:
    """A structure manifest is authoritative: its client / matter number / title are used
    verbatim, matters group under the shared client, and a '/' in a title is folder-safe."""
    source = _mini_task_set(tmp_path / "task_set")
    structure = {
        "matters": {
            "contracts/banking/account-control-agreement-first-draft/scenario-01": {
                "client": "Northwind Capital",
                "matter_no": "NC-001",
                "counterparty": "Acme",
                "matter_title": "ISDA / Credit Support Annex — Acme",
            },
            "contracts/banking/credit-support-annex-first-draft/scenario-01": {
                "client": "Northwind Capital",
                "matter_no": "NC-002",
                "counterparty": "Globex",
                "matter_title": "Account Control Agreement — Globex",
            },
        }
    }
    summary = build_task_corpus(
        tmp_path / "out", source, areas=["contracts/banking"], layout="firm", structure=structure
    )

    assert summary["clients"] == 1
    assert summary["client_names"] == ["Northwind Capital"]
    assert summary["matters"] == 2
    recs = _read_scenarios(summary)
    assert {r["matter_ref"] for r in recs} == {"NC-001", "NC-002"}
    # the '/' in the first title is sanitized to '-' so it stays one matter folder
    assert {r["matter_title"] for r in recs} == {
        "ISDA - Credit Support Annex — Acme",
        "Account Control Agreement — Globex",
    }
    acl = json.loads(Path(summary["acl_by_path"]).read_text(encoding="utf-8"))
    assert acl and all(p.startswith("Clients/Northwind Capital/NC-00") for p in acl)
    # no path segment was split by the slashed title (would appear as ".../ISDA /...")
    assert not any("/ISDA /" in p for p in acl)


def test_noise_primitives_are_deterministic_and_structural(tmp_path: Path) -> None:
    from knowledge_index.benchmark import noise

    cfg = noise.LEVELS["heavy"]
    assert noise.place("Drafts", "flat") == ""  # flat -> matter root
    assert noise.place("Drafts", "alt") == "Working Papers"  # renamed vocabulary
    assert noise.place("Drafts", "standard") == "Drafts"  # unchanged
    # a matter's style is stable across calls (seeded per matter)
    assert noise.matter_style(cfg, 42, "MCP-006") == noise.matter_style(cfg, 42, "MCP-006")

    # junk leaves an empty limbo dir and returns cruft files for the caller to wall
    matter = tmp_path / "matter"
    matter.mkdir()
    junk = noise.scatter_junk(matter, 42, "MCP-006")
    assert any(d.is_dir() and d.name.startswith("_") for d in matter.iterdir())
    assert all(j.exists() for j in junk)


def test_firm_noise_keeps_every_file_walled_and_is_deterministic(tmp_path: Path) -> None:
    """Noise may flatten/rename folders, add versions and junk — but every file stays
    walled to its own client, and the whole build is reproducible."""
    import re

    source = _mini_task_set(tmp_path / "task_set")
    summary = build_task_corpus(
        tmp_path / "out", source, areas=["contracts/banking"], layout="firm", noise_level="heavy"
    )
    assert set(summary["noise"]) == {"flat", "alt", "junk"}

    root = Path(summary["source_root"])
    acl = json.loads(Path(summary["acl_by_path"]).read_text(encoding="utf-8"))
    files = [p.relative_to(root).as_posix() for p in root.rglob("*") if p.is_file()]
    assert files and all(f in acl for f in files)  # every file (incl. junk/versions) walled
    for path, grants in acl.items():  # and walled to its OWN client — no leak
        client_slug = re.sub(r"[^a-z0-9]+", "-", path.split("/")[1].lower()).strip("-")
        assert [g["principal"] for g in grants] == [f"group:{client_slug}"]

    repeat = build_task_corpus(
        tmp_path / "out2", source, areas=["contracts/banking"], layout="firm", noise_level="heavy"
    )
    assert repeat["content_hash"] == summary["content_hash"]


def test_noise_requires_firm_layout(tmp_path: Path) -> None:
    source = _mini_task_set(tmp_path / "task_set")
    with pytest.raises(ValueError, match="firm layout"):
        build_task_corpus(
            tmp_path / "out", source, areas=["contracts/banking"], layout="flat", noise_level="light"
        )


def _read_scenarios(summary: dict) -> list[dict]:
    lines = Path(summary["scenarios_manifest"]).read_text(encoding="utf-8").splitlines()
    return [json.loads(line) for line in lines if line.strip()]


_MARKERS = ("Pinnacle Capital Partners, LLC", "Meridian Engineered Systems, Inc.")


def _marker_propose_fn(scenario: dict, doc_texts: dict[str, str], per_scenario: int) -> list:
    """Deterministic stand-in for the LLM: propose the scenario's own marker (which
    verification accepts) plus one hallucination (which it must drop)."""
    from knowledge_index.benchmark.gold import _Proposal

    marker = next(m for m in _MARKERS if any(m.casefold() in t for t in doc_texts.values()))
    return [
        _Proposal(
            query=f"Who is the secured party on the {marker.split()[0]} deal?",
            answer=marker,
            source_document="playbook.docx",
        ),
        _Proposal(
            query=f"Which entity acts as Administrative Agent for {marker.split()[0]}?",
            answer=marker,
            source_document="playbook.docx",
        ),
        _Proposal(
            query="What is the charter number?",
            answer="OCC Charter No. 99999",  # hallucinated — appears in no document
            source_document="playbook.docx",
        ),
    ]


def _generate_marker_gold(corpus_dir: Path) -> dict:
    return generate_gold(corpus_dir, propose_fn=_marker_propose_fn)


def test_generate_gold_verifies_dedupes_and_reports(tmp_path: Path) -> None:
    source = _mini_task_set(tmp_path / "task_set")
    build_task_corpus(tmp_path / "out", source, areas=["contracts/banking"])

    stats = _generate_marker_gold(tmp_path / "out")
    # 2 scenarios × 3 proposals; the hallucination is rejected per scenario
    assert stats["proposed"] == 6
    assert stats["accepted"] == 4
    assert sum(stats["by_anchor"].values()) == 4
    assert stats["rejected"] == 2
    assert stats["reject_reasons"] == {"not_in_any_readable_doc": 2}

    gold_file = tmp_path / "out" / "retrieval-gold.jsonl"
    rows = [json.loads(line) for line in gold_file.read_text().splitlines() if line.strip()]
    assert len(rows) == 4
    for row in rows:
        # gold is the verified location, and the marker lives only in the playbook
        assert len(row["gold_paths"]) == 1
        assert row["gold_paths"][0].endswith("playbook.docx")
        assert row["meta"]["answer"] in _MARKERS
        assert row["principals"]
    # rejections are diagnosed, not silently discarded
    rejected = (tmp_path / "out" / "rejected.jsonl").read_text()
    assert "not_in_any_readable_doc" in rejected

    # re-running RESUMES: scenarios already on disk are finished work and are skipped
    # entirely (not re-proposed and deduped), so a kill costs only work in flight
    again = _generate_marker_gold(tmp_path / "out")
    assert again["resumed_scenarios"] == 2
    assert again["scenarios"] == 0  # nothing left to do
    assert again["accepted"] == 0 and again["proposed"] == 0
    assert len((tmp_path / "out" / "retrieval-gold.jsonl").read_text().splitlines()) == 4


def test_stratified_sample_spreads_across_practice_areas() -> None:
    from knowledge_index.benchmark.gold import stratified_sample

    scenarios = [
        {"scenario_id": f"s{i}", "practice_area": area}
        for area in ("banking", "ip", "litigation")
        for i in range(10)
    ]
    picked = stratified_sample(scenarios, 6, seed=1)
    assert len(picked) == 6
    by_area = {area: sum(1 for s in picked if s["practice_area"] == area)
               for area in ("banking", "ip", "litigation")}
    assert by_area == {"banking": 2, "ip": 2, "litigation": 2}
    # deterministic for a given seed; no limit returns everything untouched
    assert stratified_sample(scenarios, 6, seed=1) == picked
    assert stratified_sample(scenarios, None) == scenarios


# ------------------------------------------------------------------------- gold store


def test_freeze_makes_gold_a_committed_artifact_resolvable_by_name(tmp_path: Path) -> None:
    from knowledge_index.benchmark import store

    source = _mini_task_set(tmp_path / "task_set")
    build_task_corpus(tmp_path / "out", source, areas=["contracts/banking"])
    _generate_marker_gold(tmp_path / "out")

    data_dir = tmp_path / "frozen"
    info = store.freeze(tmp_path / "out", "banking-test", data_dir=data_dir)
    assert (data_dir / "banking-test.gold.jsonl").exists()
    meta = json.loads((data_dir / "banking-test.meta.json").read_text())
    assert meta["source"].startswith("harveyai/harvey-labs")
    assert meta["corpus_config"]["content_hash"] == info["corpus_config"]["content_hash"]
    assert store.list_frozen(data_dir=data_dir) == ["banking-test"]

    by_name = store.resolve("banking-test", data_dir=data_dir)
    by_path = store.resolve(str(data_dir / "banking-test.gold.jsonl"), data_dir=data_dir)
    assert by_name == by_path
    with pytest.raises(FileNotFoundError):
        store.resolve("does-not-exist", data_dir=data_dir)


# ----------------------------------------------------------------------- measurement


def test_percentiles_and_spend_delta() -> None:
    from knowledge_index.benchmark import measure

    assert measure.percentiles([]) == {"count": 0}
    p = measure.percentiles([10.0, 20.0, 30.0, 40.0])
    assert p["count"] == 4 and p["p50"] == 30.0 and p["max"] == 40.0

    before = {"total": 0.10, "by_model": {"gpt-4.1-mini": 0.08, TEST_EMBEDDING_MODEL: 0.02}}
    after = {"total": 0.25, "by_model": {"gpt-4.1-mini": 0.20, TEST_EMBEDDING_MODEL: 0.05}}
    delta = measure.spend_delta(before, after)
    assert delta["total"] == 0.15
    assert delta["by_model"]["gpt-4.1-mini"] == 0.12

    # unavailable spend propagates rather than being estimated
    assert measure.spend_delta({"total": None, "error": "x"}, after)["total"] is None


# --------------------------------------------------------------------- harness scoring


@dataclass
class _FakeHit:
    source_paths: list[str]
    version_status: str
    excerpt: str = ""


def test_run_queries_scores_and_flags_wall_leaks() -> None:
    gold = [
        {
            "id": "M-2026-0001/first-draft#ws",
            "kind": "question",
            "query": "draft the agreement",
            "gold_paths": ["a.docx", "b.docx"],
            "principals": ["group:x-team"],
            "matter_ref": "M-2026-0001",
            "practice_area": "x",
            "meta": {"answer": "Pinnacle Capital Partners, LLC"},
        }
    ]

    def search(query, principals, filters):
        return [
            _FakeHit(["a.docx"], "final"),
            _FakeHit(["distractor.docx"], "draft"),
            _FakeHit(["b.docx"], "final", "…secured party is Pinnacle Capital\nPartners, LLC…"),
        ]

    clean = run_queries(search, gold, outsider_search_fn=lambda *a: [])
    overall = clean["metrics"]["overall"]
    assert overall["recall"]["@1"] == 0.5
    assert overall["recall"]["@5"] == 1.0
    assert overall["mrr"] == 1.0
    assert clean["observations"]["final_not_draft_rate"] == 1.0
    assert clean["ethical_wall"]["clean"] is True
    # passage-level auxiliary: the answer text only appears from rank 3 on
    # (whitespace-normalized on both sides)
    assert clean["answer_in_context"]["queries_with_answer"] == 1
    assert clean["answer_in_context"]["@1"] == 0.0
    assert clean["answer_in_context"]["@5"] == 1.0

    leak_fn = lambda *a: [_FakeHit(["a.docx"], "final")]  # noqa: E731
    leaked = run_queries(search, gold, outsider_search_fn=leak_fn)
    assert leaked["ethical_wall"]["clean"] is False
    assert leaked["ethical_wall"]["leaks"] == ["M-2026-0001/first-draft#ws"]


# ------------------------------------------------------------------- gold verification


def _proposal(**kw):
    from knowledge_index.benchmark.gold import _Proposal

    kw.setdefault("source_document", "term-sheet.docx")
    return _Proposal(**kw)


def _scenario() -> dict:
    return {
        "principal": "group:x-team",
        "matter_ref": "M-2026-0001",
        "practice_area": "contracts/banking",
        "document_paths": ["Mandate/m/term-sheet.docx", "Mandate/m/memo.docx",
                           "Mandate/m/other.docx"],
    }


def _corpus() -> dict[str, tuple[str, str]]:
    """path → (principal, text): the bundle plus scope distractors + a walled-off doc."""
    return {
        "Mandate/m/term-sheet.docx": (
            "group:x-team", "the secured party is pinnacle capital partners, llc."),
        "Mandate/m/memo.docx": ("group:x-team", "pinnacle capital partners, llc acts as agent."),
        "Mandate/m/other.docx": ("group:x-team", "unrelated content about scheduling."),
        # same value in ANOTHER team's scope: invisible to the query, must not count
        "Mandate/z/walled.docx": (
            "group:other-team", "pinnacle capital partners, llc appears here too."),
        # scope distractors that recur a generic value
        **{
            f"Mandate/d{i}/boilerplate.docx": ("group:x-team", "governed by new york law.")
            for i in range(8)
        },
    }


def test_ground_proposal_verifies_verbatim_and_widens_gold_to_the_visible_scope() -> None:
    from knowledge_index.benchmark.gold import ground_proposal

    scenario = _scenario()
    corpus = _corpus()
    bundle = {path: corpus[path][1] for path in scenario["document_paths"]}

    # grounded value in two visible docs -> gold is BOTH; the model's wrong guess and
    # the other team's walled copy are both ignored
    hit = ground_proposal(
        _proposal(query="Who is the secured party on the Pinnacle deal?",
                  answer="Pinnacle Capital Partners, LLC",
                  source_document="other.docx"),
        scenario, bundle, corpus, max_gold_docs=5, model=TEST_LLM_MODEL,
    )
    assert not isinstance(hit, dict)
    assert hit.gold_paths == ["Mandate/m/memo.docx", "Mandate/m/term-sheet.docx"]
    assert hit.kind == "request"
    assert hit.meta["extracted_by"] == f"{TEST_LLM_MODEL}/request-gold-3"
    assert hit.meta["scope_matches"] == 2
    # graded: the document the answer came from is primary, co-mentions secondary
    assert hit.meta["primary_path"] in hit.gold_paths
    assert hit.meta["secondary_paths"] == [
        p for p in hit.gold_paths if p != hit.meta["primary_path"]
    ]

    # answer that appears in no bundle document -> rejected with a diagnosis
    rejection = ground_proposal(
        _proposal(query="charter?", answer="OCC Charter No. 99999"),
        scenario, bundle, corpus, max_gold_docs=5, model=TEST_LLM_MODEL,
    )
    assert isinstance(rejection, dict)
    assert rejection["reason"] == "not_in_any_readable_doc"

    # trivially short answer -> rejected
    assert isinstance(
        ground_proposal(_proposal(query="x", answer="a"), scenario, bundle, corpus,
                        max_gold_docs=5, model=TEST_LLM_MODEL),
        dict,
    )


def test_ground_proposal_rejects_a_request_that_restates_its_own_answer() -> None:
    """A "query" identical to its answer asks a retriever to recompute the label's
    own definition — a lexical baseline satisfies it by construction."""
    from knowledge_index.benchmark.gold import ground_proposal

    scenario, corpus = _scenario(), _corpus()
    bundle = {path: corpus[path][1] for path in scenario["document_paths"]}
    rejection = ground_proposal(
        _proposal(query="Pinnacle Capital Partners, LLC",
                  answer="Pinnacle Capital Partners, LLC"),
        scenario, bundle, corpus, max_gold_docs=5, model=TEST_LLM_MODEL,
    )
    assert isinstance(rejection, dict)
    assert rejection["reason"] == "request_restates_its_own_answer"


def test_anchor_classification_slices_by_what_the_request_cites() -> None:
    from knowledge_index.benchmark.gold import classify_anchor

    assert classify_anchor("What's the status on charter 4478291?") == "identifier"
    # a request can TURN ON an identifier without containing one — the answer decides
    assert classify_anchor("What's the docket number on the Apex complaint?",
                           "25-CVS-4471") == "identifier"
    assert classify_anchor("Who signed for us?", "Hargrove & Finch LLP") == "entity"
    assert classify_anchor("What's the cap?", "$12 million") == "amount"
    assert classify_anchor("Who drafted the first markup on the deal?",
                           "outside counsel") == "none"


def test_ground_proposal_drops_non_discriminative_values() -> None:
    from knowledge_index.benchmark.gold import ground_proposal

    scenario = _scenario()
    corpus = _corpus()
    bundle = dict(corpus_bundle := {
        "Mandate/m/other.docx": "unrelated content about scheduling. governed by new york law.",
    })
    scenario = {**scenario, "document_paths": list(corpus_bundle)}
    corpus = {**corpus, "Mandate/m/other.docx": ("group:x-team", bundle["Mandate/m/other.docx"])}

    # "new york law" appears in 9 visible docs — identifies nothing, must be dropped
    rejection = ground_proposal(
        _proposal(query="Which law governs the Cascade agreement?", answer="new york law"),
        scenario, bundle, corpus, max_gold_docs=5, model=TEST_LLM_MODEL,
    )
    assert isinstance(rejection, dict)
    assert rejection["reason"] == "not_discriminative"
    assert rejection["scope_matches"] == 9


# -------------------------------------------------------------------- coverage guard


def test_coverage_is_full_only_when_every_gold_doc_is_present() -> None:
    from knowledge_index.benchmark.retrieval_eval import CorpusCoverageError, _coverage_summary

    full = _coverage_summary({"a.docx", "b.docx"}, {"a.docx", "b.docx", "extra.docx"})
    assert full["full"] is True and full["coverage"] == 1.0 and full["missing"] == []

    partial = _coverage_summary({"a.docx", "b.docx", "c.docx"}, {"a.docx"})
    assert partial["full"] is False
    assert partial["coverage"] == round(1 / 3, 4)
    assert set(partial["missing"]) == {"b.docx", "c.docx"}

    empty = _coverage_summary(set(), set())
    assert empty["full"] is False  # nothing to benchmark is also a failure

    err = CorpusCoverageError(partial)
    assert "coverage 33" in str(err)
    assert err.coverage["full"] is False


# ------------------------------------------------------------------------ agentic eval


def test_context_recall_bounds() -> None:
    from knowledge_index.benchmark.agentic_eval import context_recall

    assert context_recall({"a.docx", "x.docx"}, {"a.docx", "b.docx"}) == 0.5
    assert context_recall(set(), {"a.docx"}) == 0.0
    assert context_recall({"a.docx", "b.docx"}, {"a.docx", "b.docx"}) == 1.0
    assert context_recall({"a.docx"}, set()) == 0.0


def test_context_recall_is_measured_against_the_answering_document() -> None:
    """An agent that finds the answering document and stops is at full recall.

    Against flat gold this scored 1/3 — the gap that made context recall read
    below the judge's pass rate.
    """
    from knowledge_index.benchmark import metrics
    from knowledge_index.benchmark.agentic_eval import context_recall

    item = {
        "gold_paths": ["a.docx", "b.docx", "c.docx"],
        "meta": {"primary_path": "b.docx"},
    }
    primary, _ = metrics.graded_gold(item)
    assert context_recall({"b.docx"}, primary) == 1.0
    assert context_recall({"a.docx", "c.docx"}, primary) == 0.0


def test_apply_grading_regrades_a_checkpoint_written_before_grading_existed() -> None:
    from knowledge_index.benchmark.agentic_eval import apply_grading

    gold = [
        {"id": "q1", "gold_paths": ["a.docx", "b.docx"], "meta": {"primary_path": "b.docx"}}
    ]
    runs = [{"id": "q1", "retrieved_paths": ["b.docx"], "context_recall": 0.5}]
    apply_grading(runs, gold)
    assert runs[0]["context_recall"] == 1.0
    assert runs[0]["primary_paths"] == ["b.docx"]

    # A run with no matching gold item is left exactly as it was.
    orphan = [{"id": "gone", "retrieved_paths": [], "context_recall": 0.25}]
    apply_grading(orphan, gold)
    assert orphan[0]["context_recall"] == 0.25


def _run(anchor: str, success: bool, recall: float) -> dict:
    return {
        "kind": "request", "anchor": anchor, "success": success, "context_recall": recall,
        "gold_paths": ["a.docx", "b.docx"],
        "retrieved_paths": ["a.docx"] if recall > 0 else [],
        "tool_calls": 3, "llm_calls": 2, "wall_seconds": 1.5,
        "usage": {"total_tokens": 1000},
    }


def test_aggregate_config_splits_metrics_by_anchor() -> None:
    from knowledge_index.benchmark.agentic_eval import aggregate_config

    runs = [
        _run("none", True, 1.0),
        _run("none", False, 0.0),
        _run("identifier", True, 1.0),
        _run("identifier", True, 0.5),
    ]
    summary = aggregate_config(runs)
    assert summary["queries"] == 4
    assert summary["success_rate"] == 0.75
    assert summary["unanchored_accuracy"] == 0.5  # requests citing nothing
    assert summary["anchored_accuracy"] == 1.0  # requests citing an identifier
    assert summary["context_recall"] == 0.625
    assert summary["any_gold_surfaced"] == 0.75  # 3 of 4 runs surfaced a gold doc
    assert summary["total_tokens"] == 4000
    assert aggregate_config([]) == {"queries": 0}


def test_sample_gold_stratifies_across_anchors_and_configs_resolve() -> None:
    from knowledge_index.benchmark.agentic_eval import (
        AGENT_CONFIGS,
        DEFAULT_CONFIGS,
        _parse_verdict,
        resolve_configs,
        sample_gold,
    )

    gold = [{"id": f"q{i}", "kind": "request", "meta": {"anchor": "none"}} for i in range(10)] + [
        {"id": f"k{i}", "kind": "request", "meta": {"anchor": "identifier"}} for i in range(10)
    ]
    picked = sample_gold(gold, 6, seed=1)
    assert len(picked) == 6
    assert sum(1 for item in picked if item["meta"]["anchor"] == "none") == 3
    assert sample_gold(gold, 6, seed=1) == picked  # deterministic
    assert sample_gold(gold, None) == gold

    assert resolve_configs("default") == DEFAULT_CONFIGS
    assert set(resolve_configs("all")) == set(AGENT_CONFIGS)
    assert resolve_configs("classic_rag,agent_full_tools") == ("classic_rag", "agent_full_tools")
    with pytest.raises(KeyError):
        resolve_configs("typo")

    # every config references a real preset
    from knowledge_index.benchmark.presets import PRESETS

    assert all(spec["preset"] in PRESETS for spec in AGENT_CONFIGS.values())

    assert _parse_verdict({"verdict": "correct"}) is True
    assert _parse_verdict({"verdict": "incorrect"}) is False
    assert _parse_verdict({}) is False


def test_tool_suite_exposes_the_real_mcp_surface() -> None:
    import asyncio
    from types import SimpleNamespace

    from sqlalchemy import create_engine
    from sqlalchemy.orm import sessionmaker

    from knowledge_index.benchmark.agent import ToolSuite
    from knowledge_index.mcp_server import create_mcp_server

    engine = create_engine("sqlite://")
    service = SimpleNamespace(session=sessionmaker(engine)(), config=AppConfig())
    suite = ToolSuite(service, "group:x")

    # one source of truth: the benchmark exposes exactly the real MCP server's tools
    server = create_mcp_server(sessionmaker(engine), AppConfig)
    expected = {tool.name for tool in asyncio.run(server.list_tools())}
    names = {spec["function"]["name"] for spec in suite.specs()}
    assert names == expected
    assert {"search_semantic", "get_document", "list_matters", "resolve_entity"} <= names
    for spec in suite.specs():
        assert spec["type"] == "function" and "parameters" in spec["function"]

    # the allowlist is the agentic matrix's tool ablation: exact surface, typos fail loud
    restricted = ToolSuite(
        SimpleNamespace(session=sessionmaker(engine)(), config=AppConfig()),
        "group:x",
        allowed_tools={"search_semantic", "get_document"},
    )
    assert {spec["function"]["name"] for spec in restricted.specs()} == {
        "search_semantic",
        "get_document",
    }
    with pytest.raises(KeyError):
        ToolSuite(
            SimpleNamespace(session=sessionmaker(engine)(), config=AppConfig()),
            "group:x",
            allowed_tools={"search_semantic", "not_a_tool"},
        )


def test_track_paths_sees_every_evidence_shape_not_just_search_hits() -> None:
    """Regression: only search hits were tracked, so get_document/resolve_entity
    results were invisible and the scorer punished configs for using them."""
    from types import SimpleNamespace

    from sqlalchemy import create_engine
    from sqlalchemy.orm import sessionmaker

    from knowledge_index.benchmark.agent import ToolSuite

    engine = create_engine("sqlite://")
    suite = ToolSuite(
        SimpleNamespace(session=sessionmaker(engine)(), config=AppConfig()),
        "group:x",
        allowed_tools={"search_semantic"},
    )
    # search-hit shape
    suite._track_paths([{"source_paths": ["Mandate/m/hit.docx"]}])
    # citation contract (get_document, resolve_entity, traverse, …)
    suite._track_paths(
        {
            "citations": [
                {
                    "source_objects": [
                        {"id": "s1", "source_id": "c1", "path": "Mandate/m/cited.docx"}
                    ]
                }
            ]
        }
    )
    # nested inside graph edges
    suite._track_paths(
        {"edges": [{"citations": [{"source_objects": [
            {"id": "s2", "external_id": "e2", "path": "Mandate/m/edge.docx"}
        ]}]}]}
    )
    assert suite.retrieved_paths == {
        "Mandate/m/hit.docx",
        "Mandate/m/cited.docx",
        "Mandate/m/edge.docx",
    }

    # a bare "path" that is not a source reference (e.g. an ontology label path)
    # must NOT count as retrieved evidence
    suite._track_paths({"matter": {"practice_area": {"path": "Contracts > Banking"}}})
    assert "Contracts > Banking" not in suite.retrieved_paths


def test_known_item_success_requires_naming_the_document() -> None:
    """Regression: membership-only scoring let 'I could not find it' answers pass
    whenever a broad search happened to include the gold document."""
    from knowledge_index.benchmark.agentic_eval import _names_a_gold_document

    gold = ["Clients/Meridian/M-001/fund-v-lpa-draft.docx"]
    assert _names_a_gold_document("It's in fund-v-lpa-draft.docx (section 4).", gold)
    # the model prose-ifies filenames; separator drift must still match
    assert _names_a_gold_document("See the Fund V LPA draft, page 12.", gold)
    assert not _names_a_gold_document("I could not find that value in the index.", gold)
    assert not _names_a_gold_document("", gold)


def test_version_status_orders_within_the_chain_not_across_documents() -> None:
    """Supersession belongs where collapse picks a version, not in the fused score."""
    from knowledge_index.retrieval import _VERSION_STATUS_ORDER

    assert (
        _VERSION_STATUS_ORDER["executed"]
        > _VERSION_STATUS_ORDER["final"]
        > _VERSION_STATUS_ORDER["unknown"]
        > _VERSION_STATUS_ORDER["draft"]
    )
    # the cross-document nudge stays gentle: at k=20 it must not override more than
    # a few pool positions (the old 1.2/0.7 range overrode 26 at k=60)
    boost = AppConfig().retrieval.version_status_boost
    assert max(boost.values()) / min(boost.values()) <= 1.15


def test_mcnemar_and_wilson_are_exact() -> None:
    assert metrics.wilson_interval(0, 0) == [0.0, 0.0]
    low, high = metrics.wilson_interval(94, 100)  # ~0.94 at n=100
    assert 0.87 < low < 0.90 and 0.96 < high < 0.98
    # a proportion at the boundary stays inside [0, 1] (where normal intervals fail)
    assert metrics.wilson_interval(100, 100)[1] == 1.0

    # 8 vs 7 discordant is a coin flip — must NOT be significant
    treatment = [True] * 8 + [False] * 7 + [True] * 50
    reference = [False] * 8 + [True] * 7 + [True] * 50
    result = metrics.mcnemar(treatment, reference)
    assert result["discordant"] == 15 and result["net"] == 1
    assert result["p_value"] == 1.0 and result["significant"] is False

    # a lopsided split is a real effect
    treatment = [True] * 25 + [False] * 2 + [True] * 50
    reference = [False] * 25 + [True] * 2 + [True] * 50
    lopsided = metrics.mcnemar(treatment, reference)
    assert lopsided["net"] == 23 and lopsided["p_value"] < 0.001
    assert lopsided["significant"] is True

    with pytest.raises(ValueError):
        metrics.mcnemar([True], [True, False])


# ---------------------------------------------------------------------- report render


def test_matrix_and_agentic_markdown_render() -> None:
    from knowledge_index.benchmark.agentic_eval import render_agentic_markdown
    from knowledge_index.benchmark.retrieval_eval import render_matrix_markdown

    def _by_kind(question: float, known_item: float) -> dict:
        return {
            "metrics": {
                "by_kind": {
                    "none": {"ndcg": {"@10": question}, "queries": 10},
                    "identifier": {"ndcg": {"@10": known_item}, "queries": 10},
                }
            }
        }

    matrix_report = {
        "presets": ["naive_dense", "full"],
        "runs": {"naive_dense": _by_kind(0.39, 0.43), "full": _by_kind(0.61, 0.63)},
        "comparison": {
            "naive_dense": {
                "ndcg@10": 0.41, "recall@10": 0.6, "mrr": 0.5, "answer_in_context@10": 0.5,
                "p95_ms": 120.0, "wall_clean": True,
                "vs_full": {"delta": -0.21, "ci95": [-0.3, -0.12], "significant": True},
            },
            "full": {"ndcg@10": 0.62, "recall@10": 0.8, "mrr": 0.7,
                     "answer_in_context@10": 0.75, "p95_ms": 150.0,
                     "wall_clean": True, "vs_full": None},
        },
        "gate": {"reference": "naive_dense", "min_lift": 0.05, "lift": 0.21,
                 "ci95": [0.12, 0.3], "significant": True, "walls_clean": True,
                 "corpus_full": True, "passed": True},
    }
    rendered = render_matrix_markdown(matrix_report)
    # kinds are reported separately — pooling hands a lexical baseline a tautological half
    assert "| full | 0.62 | 0.63 | 0.61 | — |" in rendered
    assert "| naive_dense | 0.41 | 0.43 | 0.39 |" in rendered
    assert "-0.2100 [-0.3000, -0.1200] *" in rendered
    assert "PASSED" in rendered

    agentic_report = {
        "configs": ["agent_naive", "agent_full_tools"],
        "summary": {
            "agent_naive": {"success_rate": 0.5, "success_ci95": [0.41, 0.59],
                            "anchored_accuracy": 0.4,
                            "unanchored_accuracy": 0.6, "context_recall": 0.55,
                            "avg_tool_calls": 4.0},
            "agent_full_tools": {"success_rate": 0.8, "success_ci95": [0.72, 0.86],
                                 "anchored_accuracy": 0.75,
                                 "unanchored_accuracy": 0.85, "context_recall": 0.9,
                                 "avg_tool_calls": 5.5},
        },
        "comparison_vs_reference": {
            "reference": "agent_full_tools",
            "deltas": {"agent_naive": {"n": 200, "treatment_only": 5, "reference_only": 35,
                                       "discordant": 40, "net": -30, "p_value": 0.0,
                                       "significant": True}},
        },
        "wall_probe": {"probed": 10, "leaks": [], "clean": True},
    }
    rendered = render_agentic_markdown(agentic_report)
    assert "| agent_full_tools | 0.8 [0.72, 0.86]" in rendered
    assert "-30 net (5/35 discordant), p=0.0 — SIGNIFICANT" in rendered
    assert "clean" in rendered


# ------------------------------------------------------------------------- integration


@pytest.mark.integration
def test_benchmark_runs_end_to_end_and_full_beats_naive(
    factory,
    integration_config: AppConfig,
    settle_pipeline,
    refresh_search,
    tmp_path: Path,
) -> None:
    from sqlalchemy.orm import sessionmaker

    from knowledge_index.benchmark import evaluate_matrix
    from knowledge_index.db.models import Source
    from knowledge_index.sync import LocalFilesystemSource, SyncEngine

    source = _mini_task_set(tmp_path / "task_set")
    summary = build_task_corpus(tmp_path / "out", source, areas=["contracts/banking"])
    _generate_marker_gold(tmp_path / "out")
    source_root = Path(summary["source_root"])
    acl_map = json.loads(Path(summary["acl_by_path"]).read_text(encoding="utf-8"))

    assert isinstance(factory, sessionmaker)
    with factory() as session:  # type: Session
        record = Source(
            kind="local_fs",
            display_name="task-bench",
            config={"root": str(source_root), "acl_by_path": acl_map},
        )
        session.add(record)
        session.flush()
        connector = LocalFilesystemSource(
            source_root,
            acl_resolver=lambda path: acl_map[path.relative_to(source_root).as_posix()],
        )
        SyncEngine(session, record, connector).sync()
        session.commit()

    settle_pipeline(factory, integration_config)
    refresh_search(integration_config)

    gold_file = tmp_path / "out" / "retrieval-gold.jsonl"
    report = evaluate_matrix(
        factory,
        integration_config,
        gold_file,
        presets=("naive_dense", "full"),
        min_lift=0.0,
        check_ethical_wall=True,
    )
    assert report["gate"]["walls_clean"] is True
    assert report["gate"]["corpus_full"] is True
    assert report["runs"]["full"]["corpus"]["full"] is True
    full = report["comparison"]["full"]["recall@10"]
    naive = report["comparison"]["naive_dense"]["recall@10"]
    assert full is not None and naive is not None
    assert full >= naive
    # the paired comparison is present and aligned on the same queries
    versus = report["comparison"]["naive_dense"]["vs_full"]
    assert versus is not None and versus["n"] == report["runs"]["full"]["gold_queries"]


def test_oracle_context_fills_the_answering_document_first() -> None:
    """The document that answers the request must not be the one that gets cut.

    Splitting the budget evenly across gold left the primary truncated to its cover
    page, so the oracle failed for want of context and understated the ceiling.
    """
    from knowledge_index.benchmark import agent

    captured: dict = {}

    def fake_complete(config, slot, messages, **kwargs):
        captured["prompt"] = messages[-1]["content"]
        return {"content": "ok"}

    gold_texts = {"filler.docx": "F" * 50_000, "answer.docx": "A" * 50_000}
    original = agent.gateway.complete
    agent.gateway.complete = fake_complete
    try:
        agent.run_oracle(
            "q", gold_texts, None, None, primary={"answer.docx"}, char_budget=60_000
        )
    finally:
        agent.gateway.complete = original

    prompt = captured["prompt"]
    assert prompt.index("[answer.docx]") < prompt.index("[filler.docx]")
    assert prompt.count("A") == 50_000  # the primary arrives whole, not sliced in half
    assert prompt.count("F") == 10_000  # the co-mention gets what is left, and no more


def test_traverse_once_forces_exactly_one_relation_traversal() -> None:
    """Ported from the demo harness: the prompt asks, the guard makes it happen.

    find_related_documents fired 10 times in 1,282 runs when it was only asked
    for, so the guard pins it once the model has read something — and stands
    down for a model that traversed on its own.
    """
    from knowledge_index.benchmark.agent import traverse_once

    # Nothing read yet: the model is still searching, leave it alone.
    assert traverse_once([]) is None
    assert traverse_once(["search_semantic", "search_filter"]) is None

    # Read a document and still going: pin the traversal.
    forced = traverse_once(["search_semantic", "get_document"])
    assert forced == {"type": "function", "function": {"name": "find_related_documents"}}

    # Already traversed — never force a second time.
    assert traverse_once(["search_semantic", "get_document", "find_related_documents"]) is None
    # ...including when it traversed before reading anything.
    assert traverse_once(["find_related_documents", "get_document"]) is None


def test_is_self_identifying_rejects_a_request_that_names_no_matter() -> None:
    """The prompt rule alone did not hold, so the rule is enforced by a second reader.

    The previous gold carried the same "TOO VAGUE" warning and 26% of its requests
    were still ambiguous corpus-wide, which put coin-flip noise into every config.
    """
    from knowledge_index.benchmark import gold as gold_module

    calls: list[str] = []

    def fake_complete(config, slot, messages, **kwargs):
        calls.append(messages[-1]["content"])
        verdict = "AMBIGUOUS" if "the fee letter" in messages[-1]["content"] else "SPECIFIC"
        return {"content": '{"verdict": "%s", "reason": "r"}' % verdict}

    original = gold_module.gateway.complete
    gold_module.gateway.complete = fake_complete
    try:
        assert gold_module.is_self_identifying("What's in the fee letter?", None, None)[0] is False
        assert gold_module.is_self_identifying(
            "What's the arrangement fee on the Westlake fee letter?", None, None
        )[0] is True
    finally:
        gold_module.gateway.complete = original

    # the auditor sees ONLY the request — never the matter or the documents
    assert all("REQUEST:" in c and "MATTER" not in c for c in calls)


def test_specificity_check_fails_open() -> None:
    """A gateway blip must not silently shrink the gold set."""
    from knowledge_index.benchmark import gold as gold_module

    def boom(*_a, **_k):
        raise RuntimeError("gateway down")

    original = gold_module.gateway.complete
    gold_module.gateway.complete = boom
    try:
        ok, why = gold_module.is_self_identifying("anything", None, None)
    finally:
        gold_module.gateway.complete = original
    assert ok is True
    assert why == "specificity_check_unavailable"
